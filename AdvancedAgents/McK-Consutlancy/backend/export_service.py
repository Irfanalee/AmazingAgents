import os
import uuid
from html import escape
from typing import List
from datetime import datetime

import markdown as md
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from htmldocx import HtmlToDocx
from . import excel_service

EXPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
os.makedirs(EXPORTS_DIR, exist_ok=True)


# ── DOCX ─────────────────────────────────────────────────────────────────────


def _set_cell_shading(cell, fill_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _style_docx_tables(doc):
    """Apply McKinsey-style formatting to all tables in the document."""
    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if row_idx == 0:
                    _set_cell_shading(cell, '0A2342')
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                else:
                    fill = 'EBF0FA' if row_idx % 2 == 0 else 'FFFFFF'
                    _set_cell_shading(cell, fill)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)

def generate_docx(session_name: str, analyses: List[dict]) -> str:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cover page
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("NPI Strategy Suite Report")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x0A, 0x23, 0x42)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(session_name)
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(
        f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}"
    ).font.size = Pt(11)

    doc.add_page_break()

    parser = HtmlToDocx()

    for analysis in analyses:
        title = analysis.get("title", analysis.get("prompt_id", "Analysis"))
        output = analysis.get("output", "")
        cost = analysis.get("cost_usd", 0)
        from_cache = analysis.get("from_cache", False)

        # Section heading
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x0A, 0x23, 0x42)
            run.font.size = Pt(18)

        # Cost meta line
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(
            f"Cost: ${cost:.4f}" + (" (from cache)" if from_cache else "")
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        meta_run.italic = True

        # Convert markdown → HTML → docx (handles tables, lists, bold, etc.)
        if output:
            content_html = md.markdown(
                output,
                extensions=["tables", "fenced_code"],
            )
            parser.add_html_to_document(content_html, doc)
            _style_docx_tables(doc)

        doc.add_page_break()

    filename = f"pls_report_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(EXPORTS_DIR, filename)
    doc.save(filepath)
    return filepath


# ── PDF ───────────────────────────────────────────────────────────────────────

_PDF_CSS = """
@page {
    margin: 2.5cm 2cm;
    size: A4;
}
body {
    font-family: Arial, Helvetica, sans-serif;
    color: #1a1a2e;
    line-height: 1.65;
    font-size: 11pt;
}
.cover {
    text-align: center;
    padding-top: 180px;
    page-break-after: always;
}
.cover h1 {
    font-size: 30px;
    color: #0A2342;
    margin-bottom: 14px;
}
.cover .subtitle {
    font-size: 18px;
    color: #C9A84C;
    margin-bottom: 10px;
}
.cover .date {
    font-size: 12px;
    color: #777;
    margin-top: 18px;
}
.cover .divider {
    width: 80px;
    height: 3px;
    background: #C9A84C;
    margin: 20px auto;
}
.chapter {
    page-break-before: always;
}
.chapter-title {
    font-size: 22px;
    color: #0A2342;
    font-weight: bold;
    border-bottom: 3px solid #C9A84C;
    padding-bottom: 8px;
    margin-bottom: 6px;
}
.chapter-meta {
    color: #999;
    font-size: 9pt;
    margin-bottom: 18px;
    font-style: italic;
}
h2 { font-size: 16px; color: #0A2342; margin-top: 22px; margin-bottom: 6px; }
h3 { font-size: 13px; color: #2d3748; margin-top: 16px; margin-bottom: 4px; }
p { margin: 7px 0; }
ul { margin: 8px 0 8px 22px; }
ol { margin: 8px 0 8px 22px; }
li { margin: 3px 0; }
table {
    width: 100%;
    border-collapse: collapse;
    border-spacing: 0;
    border: 1px solid #ccc;
    margin: 14px 0;
    font-size: 9pt;
    word-wrap: break-word;
}
th {
    width: 100%;
    background: #0A2342;
    color: white;
    padding: 5px 4px;
    text-align: left;
    font-weight: bold;
    border: 1px solid #ccc;
}
td {
    width: 100%;
    border: 1px solid #ddd;
    padding: 4px 4px;
}
tr:nth-child(even) td { background: #f7f7f7; }
blockquote {
    border-left: 4px solid #C9A84C;
    margin: 14px 0;
    padding: 8px 16px;
    background: #fdf8ee;
    color: #444;
}
code {
    background: #f0f0f0;
    padding: 1px 5px;
    font-size: 10px;
    font-family: monospace;
}
pre {
    background: #f0f0f0;
    padding: 12px;
    font-size: 10px;
}
strong { color: #0A2342; }
"""


def _replace_tables_with_excel_html(content_html: str, excel_tables_html: str) -> str:
    """Replace markdown-converted tables in content_html with Excel-sourced HTML tables."""
    soup = BeautifulSoup(content_html, "html.parser")
    md_tables = soup.find_all("table")
    if not md_tables:
        return excel_tables_html + str(soup)

    # Parse Excel HTML into (heading + table) units
    esoup = BeautifulSoup(excel_tables_html, "html.parser")
    units, children = [], list(esoup.children)
    i = 0
    while i < len(children):
        node = children[i]
        if hasattr(node, 'name'):
            if (node.name == 'p' and i + 1 < len(children)
                    and hasattr(children[i + 1], 'name')
                    and children[i + 1].name == 'table'):
                units.append(str(node) + str(children[i + 1]))
                i += 2
                continue
            elif node.name == 'table':
                units.append(str(node))
        i += 1

    if not units:
        return content_html

    if len(md_tables) == len(units):
        for md_tbl, unit in zip(md_tables, units):
            md_tbl.replace_with(BeautifulSoup(unit, "html.parser"))
        return str(soup)

    # Count mismatch: strip markdown tables, prepend Excel tables
    for t in md_tables:
        t.decompose()
    return excel_tables_html + str(soup)


def generate_pdf(session_name: str, analyses: List[dict]) -> str:
    from xhtml2pdf import pisa

    html_parts = [
        f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>{_PDF_CSS}</style>
</head><body>
<div class="cover">
  <h1>NPI Suite Report</h1>
  <div class="divider"></div>
  <div class="subtitle">{escape(session_name)}</div>
  <div class="date">Generated: {datetime.utcnow().strftime('%B %d, %Y')}</div>
</div>
"""
    ]

    for analysis in analyses:
        title = analysis.get("title", analysis.get("prompt_id", "Analysis"))
        output = analysis.get("output", "")
        cost = analysis.get("cost_usd", 0)
        from_cache = analysis.get("from_cache", False)

        content_html = md.markdown(
            output,
            extensions=["tables", "fenced_code"],
        )

        excel_path = analysis.get("excel_path")
        if (analysis.get("prompt_id") == "9"
                and excel_path
                and os.path.isfile(excel_path)):
            try:
                excel_html = excel_service.render_excel_tables_as_html(excel_path)
                content_html = _replace_tables_with_excel_html(content_html, excel_html)
            except Exception as e:
                print(f"[export_service] Excel table replacement failed: {e}")

        meta = f"Cost: ${cost:.4f}" + (" &mdash; served from cache" if from_cache else "")

        html_parts.append(
            f"""<div class="chapter">
  <div class="chapter-title">{escape(title)}</div>
  <div class="chapter-meta">{meta}</div>
  {content_html}
</div>
"""
        )

    html_parts.append("</body></html>")
    full_html = "\n".join(html_parts)

    filename = f"pls_report_{uuid.uuid4().hex[:8]}.pdf"
    filepath = os.path.join(EXPORTS_DIR, filename)
    with open(filepath, "wb") as f:
        result = pisa.CreatePDF(full_html.encode("utf-8"), dest=f, encoding="utf-8")

    if result.err:
        raise RuntimeError(f"PDF generation failed with {result.err} error(s)")

    return filepath
