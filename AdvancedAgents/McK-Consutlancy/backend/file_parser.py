import io
from typing import Optional

MAX_CHARS = 50_000


def parse_file(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the correct parser by extension. Returns extracted plain text."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext in ("xlsx", "xls"):
        return _parse_excel(file_bytes)
    elif ext == "docx":
        return _parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type '.{ext}'. Please upload a PDF, Excel (.xlsx/.xls), or Word (.docx) file.")


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ValueError("PDF parsing requires the 'pypdf' library. Run: pip install pypdf>=3.0.0")

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)

    full_text = "\n\n".join(pages)
    if len(full_text.strip()) < 100:
        raise ValueError(
            "This PDF appears to be a scanned image or contains no extractable text. "
            "Only text-based PDFs are supported."
        )
    return _cap(full_text)


def _parse_excel(file_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ValueError("Excel parsing requires the 'openpyxl' library.")

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sections = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Skip fully empty rows
            if not any(cell is not None for cell in row):
                continue
            rows.append("\t".join("" if cell is None else str(cell) for cell in row))
        if rows:
            sections.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
    wb.close()
    return _cap("\n\n".join(sections))


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ValueError("DOCX parsing requires the 'python-docx' library.")

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return _cap("\n".join(parts))


def _cap(text: str) -> str:
    if len(text) > MAX_CHARS:
        return text[:MAX_CHARS] + "\n\n[...document truncated at 50,000 characters...]"
    return text
